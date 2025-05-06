#
# Licensed under the MIT License. 
# Copyright (c) 2025 Perry Mertens
#
# See the LICENSE file for full license text.
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import re
from pathlib import Path
from datetime import datetime


# Risk descriptions and recommendations (OWASP API Top 10 2023)
RISK_INFO = {
    "BOLA": {
        "title": "API1:2023 – Broken Object Level Authorization",
        "description": (
            "APIs expose object IDs in their endpoints. Attackers can manipulate IDs to "
            "retrieve or modify unauthorized data."
        ),
        "recommendation": """- Implement object-level authorization checks on every request
- Use unpredictable IDs (UUID) instead of sequential integers
- Verify the requester has ownership/access rights for each object
- Centralize authorization logic
- Log and alert on failed authorization attempts"""
    },
    "BrokenAuth": {
        "title": "API2:2023 – Broken Authentication",
        "description": (
            "Incorrect authentication implementations make it possible to steal tokens "
            "or hijack sessions."
        ),
        "recommendation": """- Use MFA for sensitive actions
- Work with short-lived, cryptographically signed tokens
- Secure password/token recovery flows
- Temporarily lock accounts after too many failed attempts
- Never expose credentials in URLs or error messages"""
    },
    "Property": {
        "title": "API3:2023 – Broken Object Property Level Authorization",
        "description": (
            "Mass assignment and over-exposure: clients can see or modify fields that "
            "should remain hidden."
        ),
        "recommendation": """- Explicitly define which fields are visible/editable per role
- Validate request and response payloads with schemas
- Filter sensitive fields server-side before sending
- Use different DTOs for different access levels
- Strictly separate public and private properties"""
    },
    "Resource": {
        "title": "API4:2023 – Unrestricted Resource Consumption",
        "description": (
            "No limits on payload size, pagination or requests can lead to DoS or "
            "higher costs due to resource exhaustion."
        ),
        "recommendation": """- Implement rate limiting and quotas
- Set maximum payload sizes
- Use pagination or partial responses
- Monitor abnormal consumption
- Cache expensive operations where possible"""
    },
    "AdminAccess": {
        "title": "API5:2023 – Broken Function Level Authorization",
        "description": (
            "Complex role/function matrices often lead to endpoints being accessible to regular users "
            "when they're only meant for admins."
        ),
        "recommendation": """- Use RBAC or ABAC with deny-by-default
- Centralize authorization logic
- Thoroughly test ALL admin functions
- Require step-up authentication for critical actions
- Document and encrypt sensitive admin flows"""
    },
    "BusinessFlows": {
        "title": "API6:2023 – Unrestricted Access to Sensitive Business Flows",
        "description": (
            "Attackers can automate or abuse sensitive business processes (e.g. checkout, booking) "
            "when there are no anti-fraud measures."
        ),
        "recommendation": """- Add business context validations (e.g. balance, limits)
- Use CAPTCHA/rate limiting against bots
- Detect and block abnormal patterns
- Require step-up authentication for risky actions
- Monitor critical flows in real-time"""
    },
    "SSRF": {
        "title": "API7:2023 – Server Side Request Forgery",
        "description": (
            "APIs that fetch external URLs can be abused to access internal services or "
            "leak metadata."
        ),
        "recommendation": """- Validate & sanitize all provided URLs
- Use an allow-list of permitted domains
- Don't follow redirects or limit the number of hops
- Segment internal networks; block outgoing requests where possible
- Apply egress firewall rules"""
    },
    "Misconfig": {
        "title": "API8:2023 – Security Misconfiguration",
        "description": (
            "Default or incorrect configurations (CORS, headers, debug modes) can leak sensitive "
            "information or facilitate attacks."
        ),
        "recommendation": """- Harden systems according to security baselines
- Disable unnecessary HTTP methods
- Remove debug/test endpoints in production
- Set strict CORS policies
- Regularly review & patch configurations"""
    },
    "Inventory": {
        "title": "API9:2023 – Improper Inventory Management",
        "description": (
            "Incomplete overview of endpoints/versions leads to shadow APIs and old, "
            "unpatched versions remaining online."
        ),
        "recommendation": """- Maintain an up-to-date inventory of all endpoints
- Carefully deprecate & remove old versions
- Document each endpoint with purpose & owner
- Implement clear versioning strategy
- Proactively scan for undocumented APIs"""
    },
    "UnsafeConsumption": {
        "title": "API10:2023 – Unsafe Consumption of APIs",
        "description": (
            "Over-reliance on third-party data can lead to injections or unexpected behavior "
            "when the external API behaves unexpectedly."
        ),
        "recommendation": """- Validate & sanitize all data from third-party APIs
- Set time limits & retries
- Fail safely: handle external errors gracefully
- Keep third-party credentials secret & rotate regularly
- Continuously monitor external service behavior"""
    }
}


def create_audit_report(output_dir: Path):
    """Create comprehensive DOCX report from all scan results"""
    doc = Document()
    
    # Add title page
    doc.add_heading('API Security Audit Report. This program is created by Perry Mertens April 2025 (c)', 0)
    doc.add_paragraph('Generated on: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    paragraphs = []
    for risk in RISK_INFO.values():
        paragraphs.append(risk["title"])
    doc.add_paragraph('\n'.join(paragraphs))
    doc.add_page_break()

    # Special mapping for shorter filenames
    manual_file_map = {
        'BOLA': 'bola',
        'BrokenAuth': 'broken_auth',
        'UnsafeConsumption': 'safe_consumption', 
        # Add new mappings here if you scan for new risks
    }

    # Add content for each risk
    for risk_name, info in RISK_INFO.items():
        doc.add_heading(info["title"], level=1)

        doc.add_heading('Description', level=2)
        doc.add_paragraph(info["description"])

        doc.add_heading('Findings', level=2)
        snake = manual_file_map.get(risk_name, re.sub(r'(?<!^)(?=[A-Z])', '_', risk_name).lower())
        report_file = output_dir / f"api_{snake}_report.txt"
        if report_file.exists():
            findings = report_file.read_text(encoding='utf-8')
            clean_findings = ''.join(c for c in findings if c.isprintable() or c in '\n\r\t')
            doc.add_paragraph(clean_findings)
        else:
            doc.add_paragraph("No findings reported for this risk")

        doc.add_heading('Recommendations', level=2)
        doc.add_paragraph(info["recommendation"])

        doc.add_page_break()

    # Add summary
    doc.add_heading('Executive Summary', level=1)
    summary_file = output_dir / "api_summary_report.txt"
    if summary_file.exists():
        summary = json.loads(summary_file.read_text(encoding='utf-8'))
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1' 
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Risk'
        hdr_cells[1].text = 'OWASP Category'
        hdr_cells[2].text = 'Vulnerabilities Found'

        for risk, count in summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = risk
            row_cells[1].text = RISK_INFO.get(risk, {}).get("title", "N/A")
            row_cells[2].text = str(count)

    report_path = output_dir / "API_Security_Audit_Report.docx"
    doc.save(str(report_path))
    print(f"📄 Comprehensive report saved to: {report_path}")

def main():
    output_dir = Path(".")  # Adjust if needed
    create_audit_report(output_dir)

if __name__ == "__main__":
    main()